"""
论文生成工具 —— 支持 Markdown 和 LaTeX 两种格式
- Markdown：python-docx 排版 + Word COM 转 PDF
- LaTeX：xelatex/pdflatex 直接编译为 PDF

源文件存储：data/papers/
输出文件存储：server/static/papers/
"""

import os
import re
import uuid
import sys
import subprocess
import datetime

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from agent.tools import emit_progress

# 项目根目录
ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))

# 输出目录（PDF/DOCX）
OUTPUT_DIR = os.path.join(ROOT, "server", "static", "papers")

# 源文件目录（.md / .tex）
SOURCE_DIR = os.path.join(ROOT, "data", "papers")

# LaTeX 模板目录
TEMPLATE_DIR = os.path.join(ROOT, "tools", "latex", "templates")

SCHEMA = {
    "type": "function",
    "tag": "文档",
    "function": {
        "name": "generate_paper",
        "description": "生成学术论文 PDF。title 为标题，content 为正文（Markdown/LaTeX）。简历请用 mock_interview。",
        "parameters": {
            "type": "object",
            "properties": {
                "title": {
                    "type": "string",
                    "description": "论文标题"
                },
                "content": {
                    "type": "string",
                    "description": (
                        "论文正文内容。"
                        "Markdown 格式：使用 # 作为一级标题，## 作为二级标题，正文用普通段落。"
                        "LaTeX 格式：使用 \\section{}、\\subsection{} 等 LaTeX 命令，"
                        "公式用 $...$ 或 $$...$$，支持 \\begin{table} 等环境。"
                    )
                },
                "format": {
                    "type": "string",
                    "enum": ["markdown", "latex"],
                    "description": "论文格式：markdown（默认，简单快速）或 latex（学术排版，支持公式/表格/引用）"
                },
                "output_name": {
                    "type": "string",
                    "description": "输出文件名（不含扩展名），不填则自动生成。（可选）"
                }
            },
            "required": ["title", "content"]
        }
    }
}


# ==================== Markdown 模式 ====================

def _set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _add_cover_page(doc, title: str):
    """添加论文封面"""
    for _ in range(6):
        doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    for _ in range(4):
        doc.add_paragraph()

    info_items = [
        f"生成日期：{datetime.date.today().strftime('%Y年%m月%d日')}",
        "本文档由 AI Agent 自动生成",
    ]
    for item in info_items:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(item)
        run.font.size = Pt(12)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    doc.add_page_break()


def _setup_styles(doc: Document):
    """配置文档样式"""
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    pf.first_line_indent = Cm(0.74)

    for level, (size, name) in enumerate([(16, "Heading 1"), (14, "Heading 2"), (13, "Heading 3")], 1):
        try:
            h_style = doc.styles[name]
            h_style.font.name = "黑体"
            h_style.font.size = Pt(size)
            h_style.font.bold = True
            h_style.font.color.rgb = RGBColor(0, 0, 0)
            h_style.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            h_style.paragraph_format.space_before = Pt(12)
            h_style.paragraph_format.space_after = Pt(6)
            h_style.paragraph_format.first_line_indent = Cm(0)
        except KeyError:
            pass


def _configure_page(doc: Document):
    """配置页面设置"""
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)


def _add_header_footer(doc: Document, title: str):
    """添加页眉页脚"""
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        header_para = header.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_para.add_run(title)
        run.font.size = Pt(9)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        footer = section.footer
        footer.is_linked_to_previous = False
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run = footer_para.add_run()
        run.font.size = Pt(9)
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)


def _parse_markdown_to_docx(doc: Document, content: str):
    """将 Markdown 内容解析并写入 docx"""
    lines = content.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]

        if not line.strip():
            i += 1
            continue

        if line.startswith("# ") and not line.startswith("## "):
            text = line[2:].strip()
            para = doc.add_paragraph()
            para.style = doc.styles["Heading 1"]
            run = para.add_run(text)
            run.font.name = "黑体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            i += 1
            continue

        if line.startswith("## "):
            text = line[3:].strip()
            para = doc.add_paragraph()
            para.style = doc.styles["Heading 2"]
            run = para.add_run(text)
            run.font.name = "黑体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            i += 1
            continue

        if line.startswith("### "):
            text = line[4:].strip()
            para = doc.add_paragraph()
            para.style = doc.styles["Heading 3"]
            run = para.add_run(text)
            run.font.name = "黑体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            i += 1
            continue

        if re.match(r"^\s*[-*+]\s+", line):
            para = doc.add_paragraph(style="List Bullet")
            text = re.sub(r"^\s*[-*+]\s+", "", line)
            _add_formatted_text(para, text)
            i += 1
            while i < len(lines) and lines[i].startswith("  ") and not re.match(r"^\s*[-*+]\s+", lines[i]):
                para.add_run(" " + lines[i].strip())
                i += 1
            continue

        if re.match(r"^\s*\d+[\.\)]\s+", line):
            para = doc.add_paragraph(style="List Number")
            text = re.sub(r"^\s*\d+[\.\)]\s+", "", line)
            _add_formatted_text(para, text)
            i += 1
            continue

        if re.match(r"^[-*_]{3,}\s*$", line):
            doc.add_paragraph()
            i += 1
            continue

        para = doc.add_paragraph()
        _add_formatted_text(para, line.strip())
        i += 1


def _add_formatted_text(para, text: str):
    """添加带格式的文本（处理粗体、斜体等）"""
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = para.add_run(part[1:-1])
            run.italic = True
        else:
            run = para.add_run(part)


def _convert_to_pdf(docx_path: str, pdf_path: str) -> bool:
    """通过 PowerShell 调用 Word COM 将 docx 转 PDF"""
    ps_script = f'''
$word = New-Object -ComObject Word.Application
$word.Visible = $false
$doc = $word.Documents.Open("{docx_path}")
$doc.SaveAs([ref] "{pdf_path}", [ref] 17)
$doc.Close()
$word.Quit()
[System.Runtime.Interopservices.Marshal]::ReleaseComObject($word) | Out-Null
'''
    try:
        result = subprocess.run(
            ["powershell", "-NoProfile", "-Command", ps_script],
            capture_output=True, text=True, timeout=60
        )
        if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
            return True
        return False
    except Exception:
        return False


# ==================== LaTeX 模式 ====================

def _get_latex_engine() -> str:
    """获取 LaTeX 编译器路径"""
    tools_latex = os.path.join(ROOT, "tools", "latex")
    if tools_latex not in sys.path:
        sys.path.insert(0, tools_latex)
    from compiler import get_engine
    return get_engine()


def _has_latex() -> bool:
    """检查是否有 LaTeX 引擎"""
    return bool(_get_latex_engine())


def _generate_latex(title: str, content: str, output_name: str) -> str:
    """
    使用 LaTeX 模板生成论文
    返回: 结果消息字符串
    """
    engine = _get_latex_engine()
    if not engine:
        return (
            "未找到 LaTeX 引擎（xelatex/pdflatex）。\n\n"
            "请将 MiKTeX 便携版解压到 tools/latex/miktex/ 目录下。\n"
            "下载地址：https://miktex.org/download（选择 Portable Edition）\n\n"
            "已自动降级为 Markdown 格式，请重试。"
        )

    engine_name = os.path.basename(engine).replace(".exe", "")

    # 选择模板
    emit_progress("generate_paper", 20, "正在填充 LaTeX 模板...")
    template_path = os.path.join(TEMPLATE_DIR, "article_cn.tex")
    if not os.path.exists(template_path):
        return "LaTeX 模板文件不存在，请联系管理员。"
    with open(template_path, "r", encoding="utf-8") as f:
        template = f.read()

    # 填充模板
    tex_content = template.replace("__TITLE__", title)
    tex_content = tex_content.replace("__AUTHOR__", "AI Agent")
    tex_content = tex_content.replace("__CONTENT__", content)

    # 保存 .tex 源文件
    tex_path = os.path.join(SOURCE_DIR, f"{output_name}.tex")
    with open(tex_path, "w", encoding="utf-8") as f:
        f.write(tex_content)

    # 编译
    emit_progress("generate_paper", 45, "LaTeX 编译中（第 1 遍）...")
    from compiler import compile as latex_compile
    success, pdf_path, log = latex_compile(
        tex_path,
        output_dir=OUTPUT_DIR,
        engine=engine,
        runs=2
    )

    if not success:
        return (
            f"LaTeX 编译失败。\n\n"
            f"错误信息：\n{log}\n\n"
            f"源文件已保存到 data/papers/{output_name}.tex，可手动修改后重试。"
        )

    emit_progress("generate_paper", 100, "论文生成完成")
    return (
        f"论文已生成！（LaTeX 格式，{engine_name} 编译）\n\n"
        f"[PAPER:/static/papers/{output_name}.pdf]\n\n"
        f"LaTeX 源文件：data/papers/{output_name}.tex"
    )


# ==================== 主入口 ====================

def execute(arguments: dict) -> str:
    title = arguments.get("title", "").strip()
    content = arguments.get("content", "").strip()
    output_name = arguments.get("output_name", "").strip()
    fmt = arguments.get("format", "markdown").strip().lower()

    if not title:
        return "请提供论文标题（title 参数）。"
    if not content:
        return "请提供论文内容（content 参数）。"

    if not output_name:
        safe_title = re.sub(r'[\\/*?:"<>|]', "", title)[:30]
        import uuid
        output_name = f"{safe_title}_{datetime.date.today().strftime('%Y%m%d')}_{uuid.uuid4().hex[:6]}"

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    os.makedirs(SOURCE_DIR, exist_ok=True)

    # ===== LaTeX 模式 =====
    if fmt == "latex":
        if not _has_latex():
            # 降级为 Markdown
            fmt = "markdown"
        else:
            try:
                return _generate_latex(title, content, output_name)
            except Exception as e:
                return f"LaTeX 生成失败：{e}"

    # ===== Markdown 模式（默认） =====
    try:
        emit_progress("generate_paper", 15, "正在解析 Markdown 内容...")
        # 保存源文件到 data/papers/
        md_path = os.path.join(SOURCE_DIR, f"{output_name}.md")
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(f"# {title}\n\n{content}")

        emit_progress("generate_paper", 40, "正在排版 DOCX 文档...")
        doc = Document()
        _configure_page(doc)
        _setup_styles(doc)
        _add_cover_page(doc, title)
        _parse_markdown_to_docx(doc, content)
        _add_header_footer(doc, title)

        docx_path = os.path.join(OUTPUT_DIR, f"{output_name}.docx")
        doc.save(docx_path)

        emit_progress("generate_paper", 70, "正在转换为 PDF...")
        pdf_path = os.path.join(OUTPUT_DIR, f"{output_name}.pdf")
        pdf_ok = _convert_to_pdf(docx_path, pdf_path)
        emit_progress("generate_paper", 100, "论文生成完成")

        if pdf_ok:
            return (
                f"论文已生成！\n\n"
                f"[PAPER:/static/papers/{output_name}.pdf]\n\n"
                f"可编辑文档：/static/papers/{output_name}.docx"
            )
        else:
            return (
                f"论文已生成（DOCX格式），PDF转换失败。\n"
                f"DOCX: /static/papers/{output_name}.docx\n"
                f"（请用 Word 打开后手动另存为 PDF）"
            )

    except Exception as e:
        return f"论文生成失败：{e}"


def regenerate_document(name: str, title: str, content: str, fmt: str = "markdown") -> dict:
    """从编辑后的内容重新生成 PDF（供 API 调用）"""
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    os.makedirs(SOURCE_DIR, exist_ok=True)

    # ===== LaTeX 模式 =====
    if fmt == "latex" and _has_latex():
        try:
            # 更新 .tex 源文件
            tex_path = os.path.join(SOURCE_DIR, f"{name}.tex")
            with open(tex_path, "w", encoding="utf-8") as f:
                f.write(content)

            from compiler import compile as latex_compile
            success, pdf_path, log = latex_compile(
                tex_path,
                output_dir=OUTPUT_DIR,
                engine=_get_latex_engine(),
                runs=2
            )
            return {
                "ok": success,
                "pdf_path": f"/static/papers/{name}.pdf",
                "format": "latex",
                "error": log if not success else ""
            }
        except Exception as e:
            return {"ok": False, "error": str(e), "format": "latex"}

    # ===== Markdown 模式 =====
    # 更新 .md 源文件
    md_path = os.path.join(SOURCE_DIR, f"{name}.md")
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(f"# {title}\n\n{content}")

    doc = Document()
    _configure_page(doc)
    _setup_styles(doc)
    _add_cover_page(doc, title)
    _parse_markdown_to_docx(doc, content)
    _add_header_footer(doc, title)

    docx_path = os.path.join(OUTPUT_DIR, f"{name}.docx")
    doc.save(docx_path)

    pdf_path = os.path.join(OUTPUT_DIR, f"{name}.pdf")
    pdf_ok = _convert_to_pdf(docx_path, pdf_path)

    return {
        "ok": pdf_ok,
        "pdf_path": f"/static/papers/{name}.pdf",
        "docx_path": f"/static/papers/{name}.docx",
        "format": "markdown"
    }