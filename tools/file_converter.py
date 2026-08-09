"""
文件格式转换引擎
支持文档/图片/数据/音视频/标记语言的常见格式互转。
"""

import os
import json
import csv
import io
import shutil
import subprocess
from pathlib import Path
from datetime import datetime
from typing import Optional, List, Tuple

# ===== 转换映射表 =====
# 每个源格式 → 可转换的目标格式列表
CONVERSION_MAP = {
    # 图片
    ".png":  [".jpg", ".jpeg", ".webp", ".bmp", ".gif", ".ico", ".tiff"],
    ".jpg":  [".png", ".webp", ".bmp", ".gif", ".ico", ".tiff"],
    ".jpeg": [".png", ".webp", ".bmp", ".gif", ".ico", ".tiff"],
    ".webp": [".png", ".jpg", ".jpeg", ".bmp", ".gif"],
    ".bmp":  [".png", ".jpg", ".jpeg", ".webp"],
    ".gif":  [".png", ".jpg", ".jpeg", ".webp"],
    ".ico":  [".png"],
    ".tiff": [".png", ".jpg", ".jpeg"],
    # 文档
    ".pdf":   [".docx", ".txt", ".html"],
    ".docx":  [".pdf", ".txt", ".html"],
    ".txt":   [".pdf", ".docx", ".html"],
    ".md":    [".html", ".pdf", ".docx", ".txt"],
    ".html":  [".pdf", ".docx", ".txt", ".md"],
    ".htm":   [".pdf", ".docx", ".txt"],
    # 数据
    ".json":  [".csv", ".xml", ".xlsx", ".yaml"],
    ".csv":   [".json", ".xml", ".xlsx"],
    ".xml":   [".json", ".csv", ".yaml"],
    ".yaml":  [".json"],
    ".yml":   [".json"],
    ".xlsx":  [".csv", ".json"],
    # 音视频
    ".mp4":   [".mp3", ".m4a", ".wav", ".ogg", ".flac"],
    ".avi":   [".mp3", ".m4a", ".wav"],
    ".mkv":   [".mp3", ".m4a", ".wav"],
    ".mov":   [".mp3", ".m4a", ".wav"],
    ".webm":  [".mp3", ".wav"],
    ".mp3":   [".wav", ".m4a", ".ogg"],
    ".wav":   [".mp3", ".m4a", ".ogg"],
    ".m4a":   [".mp3", ".wav"],
    ".flac":  [".mp3", ".wav"],
    ".ogg":   [".mp3", ".wav"],
    # 代码
    ".py":   [".txt", ".html"],
    ".js":   [".txt"],
    ".ts":   [".txt"],
    ".java": [".txt"],
    ".cpp":  [".txt"],
    ".c":    [".txt"],
    ".sh":   [".txt"],
    ".bat":  [".txt"],
    ".ps1":  [".txt"],
}

# MIME 类型 → 扩展名映射
MIME_EXT_MAP = {
    "image/png": ".png", "image/jpeg": ".jpg", "image/webp": ".webp",
    "image/bmp": ".bmp", "image/gif": ".gif", "image/x-icon": ".ico",
    "image/tiff": ".tiff",
    "application/pdf": ".pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
    "text/plain": ".txt", "text/markdown": ".md", "text/html": ".html",
    "application/json": ".json", "text/csv": ".csv", "text/xml": ".xml",
    "application/xml": ".xml", "text/yaml": ".yaml",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": ".xlsx",
    "video/mp4": ".mp4", "video/x-msvideo": ".avi", "video/x-matroska": ".mkv",
    "video/quicktime": ".mov", "video/webm": ".webm",
    "audio/mpeg": ".mp3", "audio/wav": ".wav", "audio/wave": ".wav",
    "audio/x-m4a": ".m4a", "audio/flac": ".flac", "audio/ogg": ".ogg",
    "text/x-python": ".py", "application/javascript": ".js",
    "text/javascript": ".js", "text/x-java": ".java",
}

OUTPUT_BASE = Path(__file__).parent.parent / "converted_files"


def detect_format(filepath: str) -> dict:
    """识别文件格式，返回格式信息和可转换的目标列表。"""
    path = Path(filepath)
    if not path.exists():
        return {"error": f"文件不存在: {filepath}"}
    ext = path.suffix.lower()
    if ext not in CONVERSION_MAP:
        return {"error": f"不支持的文件格式: {ext}", "format": ext, "targets": []}
    st = path.stat()
    return {
        "filename": path.name,
        "format": ext,
        "size": st.st_size,
        "size_human": _human_size(st.st_size),
        "targets": CONVERSION_MAP.get(ext, []),
    }


def detect_format_from_ext(ext: str) -> dict:
    """根据扩展名返回可转换列表（用于拖拽识别）。"""
    ext = ext.lower()
    if not ext.startswith("."):
        ext = f".{ext}"
    targets = CONVERSION_MAP.get(ext, [])
    return {"format": ext, "targets": targets, "supported": len(targets) > 0}


def convert(filepath: str, target_fmt: str, output_dir: Optional[str] = None,
            progress_callback=None) -> dict:
    """
    执行格式转换。
    Args:
        filepath: 源文件路径
        target_fmt: 目标格式（含点号，如 ".pdf"）
        output_dir: 输出目录，默认自动创建
        progress_callback: 进度回调 fn(percent, message)
    Returns:
        {"success": True/False, "output": "path/to/output", "error": "..."}
    """
    source = Path(filepath)
    if not source.exists():
        return {"success": False, "error": f"文件不存在: {filepath}"}

    src_ext = source.suffix.lower()
    tgt_ext = target_fmt.lower() if target_fmt.startswith(".") else f".{target_fmt}"

    # 验证转换是否支持
    allowed = CONVERSION_MAP.get(src_ext, [])
    if tgt_ext not in allowed:
        return {"success": False, "error": f"不支持 {src_ext} → {tgt_ext} 的转换"}

    # 输出目录
    today = datetime.now().strftime("%Y-%m-%d")
    out_dir = Path(output_dir) if output_dir else (OUTPUT_BASE / today)
    out_dir.mkdir(parents=True, exist_ok=True)
    out_name = source.stem + tgt_ext
    out_path = out_dir / out_name
    # 防止覆盖
    counter = 1
    while out_path.exists():
        out_path = out_dir / f"{source.stem}_{counter}{tgt_ext}"
        counter += 1

    def _progress(pct, msg=""):
        if progress_callback:
            progress_callback(pct, msg)

    _progress(10, "准备转换...")

    try:
        # ── 图片转换 ──
        if src_ext in {".png",".jpg",".jpeg",".webp",".bmp",".gif",".ico",".tiff"} \
           and tgt_ext in {".png",".jpg",".jpeg",".webp",".bmp",".gif",".ico",".tiff"}:
            _progress(30, "读取图片...")
            from PIL import Image
            img = Image.open(source)
            # GIF 动图转静态格式时提示用户
            gif_warning = ""
            if src_ext == ".gif" and tgt_ext not in (".gif",):
                try:
                    if getattr(img, "n_frames", 1) > 1:
                        gif_warning = "（注意：GIF 动图仅转换了第一帧）"
                except Exception:
                    pass
            _progress(60, "转换中...")
            save_fmt = tgt_ext.lstrip(".").upper()
            if save_fmt == "JPG":
                save_fmt = "JPEG"
            if img.mode in ("RGBA", "P") and tgt_ext in (".jpg", ".jpeg", ".bmp"):
                img = img.convert("RGB")
            img.save(out_path, format=save_fmt)
            _progress(100, "完成")
            result = {"success": True, "output": str(out_path)}
            if gif_warning:
                result["warning"] = gif_warning
            return result

        # ── PDF → DOCX/TXT/HTML ──
        if src_ext == ".pdf":
            if tgt_ext == ".txt":
                _progress(30, "提取文本...")
                from PyPDF2 import PdfReader
                reader = PdfReader(str(source))
                text = "\n".join(p.extract_text() or "" for p in reader.pages)
                out_path.write_text(text, encoding="utf-8")
                _progress(100, "完成")
                return {"success": True, "output": str(out_path)}
            elif tgt_ext in (".docx", ".html"):
                _progress(30, "提取内容...")
                from PyPDF2 import PdfReader
                reader = PdfReader(str(source))
                text = "\n\n".join(p.extract_text() or "" for p in reader.pages)
                if tgt_ext == ".html":
                    html = f"<html><body><pre>{text}</pre></body></html>"
                    out_path.write_text(html, encoding="utf-8")
                else:
                    from docx import Document
                    doc = Document()
                    doc.add_paragraph(text)
                    doc.save(str(out_path))
                _progress(100, "完成")
                return {"success": True, "output": str(out_path)}

        # ── DOCX → PDF/TXT/HTML ──
        if src_ext == ".docx":
            from docx import Document
            doc = Document(str(source))
            text = "\n".join(p.text for p in doc.paragraphs)
            if tgt_ext == ".txt":
                out_path.write_text(text, encoding="utf-8")
            elif tgt_ext == ".html":
                html = "<html><body>\n" + "\n".join(f"<p>{p.text}</p>" for p in doc.paragraphs if p.text) + "\n</body></html>"
                out_path.write_text(html, encoding="utf-8")
            elif tgt_ext == ".pdf":
                _progress(30, "生成 PDF...")
                from reportlab.lib.pagesizes import A4
                from reportlab.platypus import SimpleDocTemplate, Paragraph
                from reportlab.lib.styles import getSampleStyleSheet
                styles = getSampleStyleSheet()
                doc_pdf = SimpleDocTemplate(str(out_path), pagesize=A4)
                flowables = [Paragraph(t.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;"), styles["Normal"])
                           for t in text.split("\n") if t]
                doc_pdf.build(flowables)
                _progress(100, "完成")
                return {"success": True, "output": str(out_path)}
            _progress(100, "完成")
            return {"success": True, "output": str(out_path)}

        # ── TXT → PDF/DOCX/HTML ──
        if src_ext == ".txt" and tgt_ext in (".docx", ".html", ".pdf"):
            text = source.read_text(encoding="utf-8", errors="replace")
            if tgt_ext == ".html":
                html = f"<html><body><pre>{text}</pre></body></html>"
                out_path.write_text(html, encoding="utf-8")
            elif tgt_ext == ".docx":
                from docx import Document
                doc = Document()
                doc.add_paragraph(text)
                doc.save(str(out_path))
            elif tgt_ext == ".pdf":
                from reportlab.lib.pagesizes import A4
                from reportlab.platypus import SimpleDocTemplate, Paragraph, PageBreak
                from reportlab.lib.styles import getSampleStyleSheet
                styles = getSampleStyleSheet()
                doc_pdf = SimpleDocTemplate(str(out_path), pagesize=A4)
                flowables = []
                for t in text.split("\n"):
                    if t.strip():
                        flowables.append(Paragraph(t.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;"), styles["Normal"]))
                doc_pdf.build(flowables)
            _progress(100, "完成")
            return {"success": True, "output": str(out_path)}

        # ── Markdown → HTML/PDF/DOCX/TXT ──
        if src_ext == ".md":
            md_text = source.read_text(encoding="utf-8", errors="replace")
            if tgt_ext == ".html":
                import markdown as md_mod
                html = md_mod.markdown(md_text, extensions=['fenced_code', 'tables'])
                out_path.write_text(html, encoding="utf-8")
            elif tgt_ext == ".txt":
                out_path.write_text(md_text, encoding="utf-8")
            elif tgt_ext == ".docx":
                from docx import Document
                import markdown as md_mod
                doc = Document()
                doc.add_paragraph(md_text)
                doc.save(str(out_path))
            elif tgt_ext == ".pdf":
                from reportlab.lib.pagesizes import A4
                from reportlab.platypus import SimpleDocTemplate, Paragraph
                from reportlab.lib.styles import getSampleStyleSheet
                styles = getSampleStyleSheet()
                doc_pdf = SimpleDocTemplate(str(out_path), pagesize=A4)
                flowables = [Paragraph(t.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;"), styles["Normal"])
                           for t in md_text.split("\n") if t.strip()]
                doc_pdf.build(flowables)
            _progress(100, "完成")
            return {"success": True, "output": str(out_path)}

        # ── HTML → PDF/DOCX/TXT ──
        if src_ext in (".html", ".htm"):
            text = source.read_text(encoding="utf-8", errors="replace")
            # Strip HTML tags for TXT
            import re
            plain = re.sub(r"<[^>]+>", " ", text)
            plain = re.sub(r"\s+", " ", plain).strip()
            if tgt_ext == ".txt":
                out_path.write_text(plain, encoding="utf-8")
            elif tgt_ext == ".docx":
                from docx import Document
                doc = Document()
                doc.add_paragraph(plain)
                doc.save(str(out_path))
            elif tgt_ext == ".pdf":
                from reportlab.lib.pagesizes import A4
                from reportlab.platypus import SimpleDocTemplate, Paragraph
                from reportlab.lib.styles import getSampleStyleSheet
                styles = getSampleStyleSheet()
                doc_pdf = SimpleDocTemplate(str(out_path), pagesize=A4)
                flowables = [Paragraph(t.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;"), styles["Normal"])
                           for t in plain.split(".\n") if t.strip()]
                if not flowables:
                    flowables = [Paragraph(plain, styles["Normal"])]
                doc_pdf.build(flowables)
            _progress(100, "完成")
            return {"success": True, "output": str(out_path)}

        # ── JSON ↔ CSV/XML/XLSX/YAML ──
        if src_ext == ".json":
            data = json.loads(source.read_text(encoding="utf-8"))
            if tgt_ext == ".csv":
                _json_to_csv(data, out_path)
            elif tgt_ext == ".xml":
                _json_to_xml(data, out_path)
            elif tgt_ext == ".xlsx":
                _json_to_xlsx(data, out_path)
            elif tgt_ext == ".yaml":
                import yaml
                out_path.write_text(yaml.dump(data, allow_unicode=True), encoding="utf-8")
            _progress(100, "完成")
            return {"success": True, "output": str(out_path)}

        if src_ext == ".csv":
            import pandas as pd
            df = pd.read_csv(source)
            if tgt_ext == ".json":
                df.to_json(out_path, orient="records", force_ascii=False, indent=2)
            elif tgt_ext == ".xml":
                df.to_xml(out_path)
            elif tgt_ext == ".xlsx":
                df.to_excel(out_path, index=False)
            _progress(100, "完成")
            return {"success": True, "output": str(out_path)}

        if src_ext == ".xml":
            import pandas as pd
            df = pd.read_xml(source)
            if tgt_ext == ".json":
                df.to_json(out_path, orient="records", force_ascii=False, indent=2)
            elif tgt_ext == ".csv":
                df.to_csv(out_path, index=False)
            elif tgt_ext == ".yaml":
                import yaml
                yaml.dump(df.to_dict(orient="records"), out_path.open("w",encoding="utf-8"), allow_unicode=True)
            _progress(100, "完成")
            return {"success": True, "output": str(out_path)}

        if src_ext == ".xlsx":
            import pandas as pd
            df = pd.read_excel(source)
            if tgt_ext == ".csv":
                df.to_csv(out_path, index=False)
            elif tgt_ext == ".json":
                df.to_json(out_path, orient="records", force_ascii=False, indent=2)
            _progress(100, "完成")
            return {"success": True, "output": str(out_path)}

        # ── 音视频转换 (ffmpeg) ──
        if src_ext in {".mp4",".avi",".mkv",".mov",".webm",".mp3",".wav",".m4a",".flac",".ogg"} \
           and tgt_ext in {".mp3",".m4a",".wav",".ogg",".flac"}:
            _progress(20, "调用 ffmpeg...")
            ffmpeg = shutil.which("ffmpeg")
            if not ffmpeg:
                # 从 config.yaml 读取 ffmpeg 路径（避免硬编码）
                ffmpeg_paths = ["ffmpeg", "ffmpeg.exe"]
                try:
                    from agent.config import get as _cfg_get
                    cfg_paths = _cfg_get("external.ffmpeg_paths")
                    if cfg_paths:
                        ffmpeg_paths.extend(cfg_paths)
                except Exception:
                    # 回退到常见路径
                    ffmpeg_paths.extend([
                        r"C:\ffmpeg\bin\ffmpeg.exe",
                        r"D:\ffmpeg\bin\ffmpeg.exe",
                    ])
                for fp in ffmpeg_paths:
                    if shutil.which(fp) or os.path.exists(fp):
                        ffmpeg = fp if os.path.exists(fp) else shutil.which(fp)
                        break
            if not ffmpeg:
                return {"success": False, "error": "未找到 ffmpeg，请安装后重试"}
            cmd = [ffmpeg, "-y", "-i", str(source)]
            if tgt_ext == ".mp3":
                cmd += ["-vn", "-acodec", "libmp3lame", "-q:a", "2"]
            elif tgt_ext == ".m4a":
                cmd += ["-vn", "-acodec", "aac", "-b:a", "192k"]
            elif tgt_ext == ".wav":
                cmd += ["-vn", "-acodec", "pcm_s16le"]
            elif tgt_ext == ".ogg":
                cmd += ["-vn", "-acodec", "libvorbis", "-q:a", "4"]
            elif tgt_ext == ".flac":
                cmd += ["-vn", "-acodec", "flac"]
            cmd.append(str(out_path))
            _progress(40, "转换中...")
            try:
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=300)
                if result.returncode != 0:
                    return {"success": False, "error": f"ffmpeg 错误: {result.stderr[:500]}"}
            except subprocess.TimeoutExpired:
                return {"success": False, "error": "ffmpeg 转换超时（5分钟），请检查文件是否损坏或过大"}
            _progress(100, "完成")
            return {"success": True, "output": str(out_path)}

        # ── 纯文本代码 → TXT/HTML ──
        if src_ext in {".py",".js",".ts",".java",".cpp",".c",".sh",".bat",".ps1"}:
            import html as _html_mod
            text = source.read_text(encoding="utf-8", errors="replace")
            if tgt_ext == ".txt":
                out_path.write_text(text, encoding="utf-8")
            elif tgt_ext == ".html":
                escaped = _html_mod.escape(text)
                html = f"<html><head><style>pre{{background:#1e1e1e;color:#d4d4d4;padding:16px;border-radius:8px;overflow:auto}}</style></head><body><pre>{escaped}</pre></body></html>"
                out_path.write_text(html, encoding="utf-8")
            _progress(100, "完成")
            return {"success": True, "output": str(out_path)}

    except Exception as e:
        return {"success": False, "error": f"{type(e).__name__}: {e}"}

    return {"success": False, "error": f"转换 {src_ext} → {tgt_ext} 暂未实现"}


# ===== 辅助函数 =====
def _human_size(size: int) -> str:
    for unit in ("B","KB","MB","GB"):
        if size < 1024:
            return f"{size:.1f} {unit}"
        size /= 1024
    return f"{size:.1f} TB"

def _json_to_csv(data, out_path):
    if isinstance(data, list):
        if not data:
            out_path.write_text("")
            return
        # 过滤非 dict 元素（混合类型列表中的纯量值）
        dict_items = [d for d in data if isinstance(d, dict)]
        if not dict_items:
            out_path.write_text("")
            return
        keys = list(dict_items[0].keys())
        writer = csv.DictWriter(out_path.open("w",encoding="utf-8",newline=""), fieldnames=keys)
        writer.writeheader()
        writer.writerows(dict_items)
    elif isinstance(data, dict):
        writer = csv.writer(out_path.open("w",encoding="utf-8",newline=""))
        writer.writerow(["key", "value"])
        for k, v in data.items():
            writer.writerow([k, str(v)])

def _json_to_xml(data, out_path):
    import re as _re
    def _sanitize_tag(name: str) -> str:
        """将非法 XML tag 名称转换为合法格式"""
        name = str(name)
        # 去掉或替换非法字符：不以数字开头，不含空格和特殊字符
        name = _re.sub(r'[^a-zA-Z0-9_\-.]', '_', name)
        if name and name[0].isdigit():
            name = '_' + name
        return name or 'item'

    def _to_xml(obj, root="root"):
        from xml.etree.ElementTree import Element, SubElement, tostring
        def _build(parent, data):
            if isinstance(data, dict):
                for k, v in data.items():
                    child = SubElement(parent, _sanitize_tag(k))
                    _build(child, v)
            elif isinstance(data, list):
                for item in data:
                    child = SubElement(parent, "item")
                    _build(child, item)
            else:
                parent.text = str(data)
        el = Element(root)
        _build(el, data)
        out_path.write_bytes(tostring(el, encoding="utf-8", xml_declaration=True))

def _json_to_xlsx(data, out_path):
    import pandas as pd
    if isinstance(data, list):
        try:
            pd.DataFrame(data).to_excel(out_path, index=False)
        except (ValueError, TypeError) as e:
            # 嵌套结构展平处理
            import json as _j
            flat_data = [{"key": str(i), "value": _j.dumps(item, ensure_ascii=False) if isinstance(item, (dict, list)) else str(item)} for i, item in enumerate(data)]
            pd.DataFrame(flat_data).to_excel(out_path, index=False)
    elif isinstance(data, dict):
        try:
            pd.DataFrame(list(data.items()), columns=["key","value"]).to_excel(out_path, index=False)
        except (ValueError, TypeError):
            import json as _j
            flat_items = [[k, _j.dumps(v, ensure_ascii=False) if isinstance(v, (dict, list)) else str(v)] for k, v in data.items()]
            pd.DataFrame(flat_items, columns=["key","value"]).to_excel(out_path, index=False)
